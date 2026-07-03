import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ===== 页面设置 =====
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ===== 样式辅助函数 =====
def set_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_bottom_border(paragraph, color='1a73e8'):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_bullet(doc, text, bold_prefix=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run('• ')
    set_font(run, size=10.5, color=(0x33, 0x33, 0x33))
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_font(run, size=10.5, bold=True, color=(0x33, 0x33, 0x33))
    run = p.add_run(text)
    set_font(run, size=10.5, color=(0x33, 0x33, 0x33))
    return p

def add_section_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    set_font(run, size=13, bold=True, color=(0x1a, 0x73, 0xe8))
    add_bottom_border(p, '1a73e8')
    return p

# ===== 顶部：个人信息行 =====
from docx.shared import Inches, Pt as PtShared, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.paragraph_format.space_before = Pt(6)
info_p.paragraph_format.space_after = Pt(8)

run = info_p.add_run('姓名：代洋')
set_font(run, size=10.5, color=(0x33, 0x33, 0x33))
run = info_p.add_run('    |    ')
set_font(run, size=10.5, color=(0x99, 0x99, 0x99))
run = info_p.add_run('学校：成都工业学院')
set_font(run, size=10.5, color=(0x33, 0x33, 0x33))
run = info_p.add_run('    |    ')
set_font(run, size=10.5, color=(0x99, 0x99, 0x99))
run = info_p.add_run('专业：数据科学与大数据技术')
set_font(run, size=10.5, color=(0x33, 0x33, 0x33))
run = info_p.add_run('    |    ')
set_font(run, size=10.5, color=(0x99, 0x99, 0x99))
run = info_p.add_run('年级：大二（2027届）')
set_font(run, size=10.5, color=(0x33, 0x33, 0x33))

# 联系方式行
contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_before = Pt(0)
contact_p.paragraph_format.space_after = Pt(4)
run = contact_p.add_run('📧 （填写邮箱）')
set_font(run, size=9.5, color=(0x66, 0x66, 0x66))
run = contact_p.add_run('  |  ')
set_font(run, size=9.5, color=(0xbb, 0xbb, 0xbb))
run = contact_p.add_run('📱 （填写电话）')
set_font(run, size=9.5, color=(0x66, 0x66, 0x66))
run = contact_p.add_run('  |  ')
set_font(run, size=9.5, color=(0xbb, 0xbb, 0xbb))
run = contact_p.add_run('📍 成都')
set_font(run, size=9.5, color=(0x66, 0x66, 0x66))

# ===== 求职意向（浅蓝背景标签） =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(6)
pPr = p._element.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'e8f0fe')
pPr.append(shd)
ind = OxmlElement('w:ind')
ind.set(qn('w:left'), '240')
ind.set(qn('w:right'), '240')
pPr.append(ind)

run = p.add_run('🎯 求职意向：AI Agent 开发实习生 ｜ LLM 应用开发实习生 ｜ Python 后端开发实习生')
set_font(run, size=10.5, bold=True, color=(0x1a, 0x1a, 0x2e))

# ===== 专业技能 =====
add_section_title(doc, '🛠 专业技能')
skills_text = [
    'AI Agent 开发：熟悉 LangGraph / LangChain，能够基于 StateGraph 设计多节点 Agent 流程、条件路由、Human-in-the-loop（interrupt/resume）、多智能体调度（Supervisor 模式）和工具调用逻辑',
    '后端开发：熟悉 Python、FastAPI、RESTful API 设计和 SSE 流式输出，能够完成 AI 应用接口设计、会话管理（SQLite 持久化）和文件上传功能',
    '大模型应用：熟悉通义千问 VL 多模态模型、Prompt Engineering、Function Calling，能够完成图片识别、结构化输出等任务',
    'RAG 技术：熟悉 ChromaDB 向量数据库、Embedding 模型、文本分块策略（RecursiveCharacterTextSplitter），有完整的 RAG 知识库构建和检索融合经验',
    '数据库与工具：MySQL、SQLite、Git、阿里云 OSS、Tavily Search API、LangSmith 调试追踪',
]
for s in skills_text:
    add_bullet(doc, s)

# ===== 项目经历 =====
add_section_title(doc, '📌 项目经历')

# 项目标题行
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('私厨助手 AI Agent（Personal Chief）')
set_font(run, size=11, bold=True, color=(0x1a, 0x1a, 0x2e))
run = p.add_run('   2026.05 - 至今')
set_font(run, size=10, color=(0x88, 0x88, 0x88))

# 技术栈
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('技术栈：LangGraph · FastAPI · ChromaDB · Tavily · DashScope VL · LangSmith')
set_font(run, size=9.5, color=(0x1a, 0x73, 0xe8))

project_items = [
    ('项目背景：', '针对用户日常饮食场景中"不知道吃什么、怎么做"的痛点，独立设计并开发了一套基于 LangGraph 的多智能体 AI 应用。支持图片识别食材、菜谱推荐、营养分析、健身饮食规划四大功能，从需求分析到上线部署全流程独立完成。'),
    ('技术架构：', '基于 LangGraph 自定义 StateGraph，设计 7 个功能节点和 3 条条件边，构建了 Supervisor + 3 个子 Agent（食谱推荐/营养分析/健身饮食）的多智能体架构。通过 LLM 意图识别实现智能路由，各 Agent 共享 State 通信，新增一个 Agent 仅需 10 分钟。'),
    ('RAG 知识库：', '构建 54 道中文菜谱的向量知识库，采用 RecursiveCharacterTextSplitter 分块（chunk_size=400, overlap=50）存入 ChromaDB，融合 Tavily 实时网页搜索实现双路检索，覆盖本地知识 + 互联网信息的互补查询。'),
    ('Human-in-the-loop：', '使用 interrupt() / Command(resume=) 实现推荐前暂停询问用户饮食偏好，用户回复后自动恢复执行流程，显著提升推荐结果的个性化程度和用户体验。'),
    ('全栈工程化：', 'FastAPI 构建 5 个 RESTful 接口；前端实现 Markdown 渲染和 SSE 流式展示；SqliteSaver 实现多会话持久化；项目已通过 Docker 部署上线（Railway），支持公网访问。'),
]
for bold_prefix, text in project_items:
    add_bullet(doc, text, bold_prefix)

# 项目成果（STAR法则 + 数据支撑）
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('项目成果（STAR）：')
set_font(run, size=10.5, bold=True, color=(0x1a, 0x73, 0xe8))

result_items = [
    '成果：从 0 到 1 独立完成全链路开发（需求→架构→编码→RAG→前端→Docker部署上线），项目已开源并部署到 Railway，支持公网访问',
    '数据：7 个功能节点 + 3 条条件边构建 Agent 工作流；54 道菜谱的向量知识库；1 个 Supervisor 调度 3 个子 Agent',
    '亮点：手写 StateGraph 替代 create_agent 黑盒，实现 HITL 人机交互暂停恢复、多智能体路由等高级特性',
]
for item in result_items:
    add_bullet(doc, item)

# ===== 荣誉奖项 =====
add_section_title(doc, '🏆 荣誉奖项')
add_bullet(doc, '蓝桥杯全国软件和信息技术专业人才大赛 省级三等奖')
add_bullet(doc, '全国大学生统计建模大赛 校级三等奖')
add_bullet(doc, '全国大学生市场调查与分析大赛 校级三等奖')

# ===== 自我评价 =====
add_section_title(doc, '💡 自我评价')
evaluations = [
    '具备扎实的 Python 编程基础和 AI Agent 开发能力，能独立完成从需求分析到开发上线的全流程',
    '对 LangGraph / LangChain 生态有深入理解，熟悉大模型应用开发范式和 RAG 技术体系',
    '自驱力强，大二期间自学 AI Agent 开发，善于通过项目驱动学习，有完整的开源项目经验',
]
for e in evaluations:
    add_bullet(doc, e)

# ===== 保存 =====
doc.save('jianli-代洋.docx')
print('✅ 简历已生成：jianli-代洋.docx')
print('请手动填写：电话和邮箱（当前为占位符）')
